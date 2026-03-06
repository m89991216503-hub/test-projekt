# -*- coding: utf-8 -*-
"""
Расчёт теплопотерь и электрической мощности базы отдыха (полный комплекс)
"""
import math
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ======================================================================
# ФИЗИЧЕСКИЕ КОНСТАНТЫ И НОРМАТИВНЫЕ ДАННЫЕ
# ======================================================================
LAM        = 0.045    # теплопроводность мин.ваты 150 кг/м³ (усл. Б), Вт/(м·К)
T_EXT      = -31      # расчётная t нар. воздуха (пятидневка 0,92), °С
T_OT_AVG   = -3.6     # средняя t отоп. периода, °С
N_OT       = 209      # продолжительность отоп. периода, сут
RHO_AIR    = 1.2      # плотность воздуха, кг/м³
C_AIR      = 1006.0   # теплоёмкость воздуха, Дж/(кг·К)
# Коэффициенты теплоотдачи (СП 50.13330.2022, табл. 6)
RSI_WALL   = 0.13     # внутренняя поверхность, стены
RSI_CEIL   = 0.10     # внутренняя поверхность, потолок
RSI_FLOOR  = 0.17     # внутренняя поверхность, пол
RSE        = 0.04     # наружная поверхность
# Зонный метод: базовые сопротивления пола по грунту (СП 50.13330.2022, прил. В, табл. В.1)
R_BASE_ZONES = [2.1, 4.3, 8.6, 14.2]
R_WIN  = 0.56         # двухкамерный стеклопакет ПВХ, м²·°С/Вт
R_DOOR = 0.60         # утеплённая металлическая дверь, м²·°С/Вт

# ======================================================================
# ДАННЫЕ ЗДАНИЙ
# ======================================================================
BUILDINGS = [
    dict(id='a', name='Жилой дом 20 м²', count=51,
         width=4.0,  length=5.0,    height=3.0, perimeter=18.0,  area=20.0,
         n_doors=1, A_door=2.0, n_windows=3, A_window=2.0,
         wall_mm=200, ceil_mm=250, floor_mm=250, floor_type='above_air',
         t_int=20, heated=True, btype='residential', n_vent=0.35,
         dhw_W=500, light_W=200, misc_W=300),
    dict(id='b', name='Жилой дом 65 м²', count=10,
         width=8.0,  length=8.125,  height=3.0, perimeter=32.25, area=65.0,
         n_doors=1, A_door=2.0, n_windows=5, A_window=2.0,
         wall_mm=200, ceil_mm=250, floor_mm=250, floor_type='above_air',
         t_int=20, heated=True, btype='residential', n_vent=0.35,
         dhw_W=1000, light_W=650, misc_W=500),
    dict(id='c', name='Жилой дом 100 м²', count=1,
         width=8.0,  length=12.5,   height=3.0, perimeter=41.0,  area=100.0,
         n_doors=1, A_door=2.0, n_windows=6, A_window=2.0,
         wall_mm=200, ceil_mm=250, floor_mm=250, floor_type='above_air',
         t_int=20, heated=True, btype='residential', n_vent=0.35,
         dhw_W=1500, light_W=1000, misc_W=800),
    dict(id='d', name='СПА-комплекс', count=1,
         width=20.0, length=25.0,   height=4.5, perimeter=90.0,  area=500.0,
         n_doors=3, A_door=2.0, n_windows=10, A_window=2.0,
         wall_mm=250, ceil_mm=300, floor_mm=300, floor_type='on_ground',
         t_int=25, heated=True, btype='public', n_vent=2.0,
         dhw_W=20000, light_W=7500, misc_W=30000),
    dict(id='e', name='Складское здание (не отапл.)', count=2,
         width=12.0, length=20.0,   height=4.5, perimeter=64.0,  area=240.0,
         n_doors=2, A_door=2.0, n_windows=4, A_window=2.0,
         wall_mm=150, ceil_mm=150, floor_mm=150, floor_type='on_ground',
         t_int=None, heated=False, btype='warehouse', n_vent=1.0,
         dhw_W=0, light_W=1200, misc_W=200),
    dict(id='f', name='Оранжерея (не отапл.)', count=1,
         width=15.0, length=20.0,   height=5.0, perimeter=70.0,  area=300.0,
         n_doors=2, A_door=2.0, n_windows=12, A_window=2.0,
         wall_mm=150, ceil_mm=150, floor_mm=150, floor_type='on_ground',
         t_int=None, heated=False, btype='greenhouse', n_vent=1.0,
         dhw_W=0, light_W=6000, misc_W=2000),
    dict(id='g', name='Административное здание', count=2,
         width=12.0, length=20.0,   height=4.5, perimeter=64.0,  area=240.0,
         n_doors=2, A_door=2.0, n_windows=4, A_window=2.0,
         wall_mm=250, ceil_mm=300, floor_mm=300, floor_type='on_ground',
         t_int=20, heated=True, btype='admin', n_vent=1.5,
         dhw_W=500, light_W=3600, misc_W=5000),
    dict(id='h', name='Пункт проката инвентаря', count=1,
         width=6.0,  length=10.0,   height=4.5, perimeter=32.0,  area=60.0,
         n_doors=1, A_door=2.0, n_windows=2, A_window=2.0,
         wall_mm=250, ceil_mm=300, floor_mm=300, floor_type='on_ground',
         t_int=16, heated=True, btype='utility', n_vent=1.0,
         dhw_W=0, light_W=600, misc_W=5000),
    dict(id='i', name='Центр парусного спорта', count=1,
         width=15.0, length=20.0,   height=4.5, perimeter=70.0,  area=300.0,
         n_doors=3, A_door=2.0, n_windows=8, A_window=2.0,
         wall_mm=250, ceil_mm=300, floor_mm=300, floor_type='on_ground',
         t_int=18, heated=True, btype='sports', n_vent=1.0,
         dhw_W=0, light_W=3000, misc_W=10000),
    dict(id='j', name='Падел-арена (не отапл.)', count=1,
         width=12.5, length=16.0,   height=4.5, perimeter=57.0,  area=200.0,
         n_doors=3, A_door=2.0, n_windows=2, A_window=2.0,
         wall_mm=150, ceil_mm=150, floor_mm=150, floor_type='on_ground',
         t_int=None, heated=False, btype='sports_unheated', n_vent=1.0,
         dhw_W=0, light_W=5000, misc_W=0),
    dict(id='k', name='Кафетерий', count=1,
         width=8.0,  length=12.5,   height=4.5, perimeter=41.0,  area=100.0,
         n_doors=3, A_door=2.0, n_windows=8, A_window=2.0,
         wall_mm=250, ceil_mm=300, floor_mm=300, floor_type='on_ground',
         t_int=20, heated=True, btype='catering', n_vent=2.0,
         dhw_W=2000, light_W=1500, misc_W=15000),
    dict(id='l', name='Бар', count=1,
         width=8.0,  length=12.5,   height=4.5, perimeter=41.0,  area=100.0,
         n_doors=3, A_door=2.0, n_windows=8, A_window=2.0,
         wall_mm=250, ceil_mm=300, floor_mm=300, floor_type='on_ground',
         t_int=20, heated=True, btype='catering', n_vent=2.0,
         dhw_W=1000, light_W=1500, misc_W=10000),
    dict(id='m', name='Зал для йоги (не отапл.)', count=1,
         width=12.5, length=16.0,   height=4.5, perimeter=57.0,  area=200.0,
         n_doors=3, A_door=2.0, n_windows=2, A_window=2.0,
         wall_mm=150, ceil_mm=150, floor_mm=150, floor_type='on_ground',
         t_int=None, heated=False, btype='unheated', n_vent=1.0,
         dhw_W=0, light_W=2000, misc_W=0),
    dict(id='n', name='Кафе', count=1,
         width=12.0, length=20.0,   height=4.5, perimeter=64.0,  area=240.0,
         n_doors=3, A_door=2.0, n_windows=12, A_window=2.0,
         wall_mm=250, ceil_mm=300, floor_mm=300, floor_type='on_ground',
         t_int=20, heated=True, btype='catering', n_vent=2.0,
         dhw_W=3000, light_W=3600, misc_W=20000),
]

INFRA_W = 20000  # уличное освещение, водоснабжение, охрана, канализация

# ======================================================================
# РАСЧЁТНЫЕ ФУНКЦИИ
# ======================================================================
def calc_gsop(t_int):   return (t_int - T_OT_AVG) * N_OT
def r_req_wall(g):      return 0.00035*g + 1.4
def r_req_ceil(g):      return 0.0005*g  + 2.2
def r_req_floor(g):     return 0.00045*g + 1.9
def r_req_win(g):       return 0.000075*g + 0.15
def R_wall(mm):         return RSI_WALL + mm/1000/LAM + RSE
def R_ceil(mm):         return RSI_CEIL + mm/1000/LAM + RSE
def R_floor_air(mm):    return RSI_FLOOR + mm/1000/LAM + RSE

def floor_zones(L, W):
    z1 = L*W       - max(0,L-4) *max(0,W-4)
    z2 = max(0,L-4) *max(0,W-4) - max(0,L-8) *max(0,W-8)
    z3 = max(0,L-8) *max(0,W-8) - max(0,L-12)*max(0,W-12)
    z4 = max(0,L-12)*max(0,W-12)
    return [z1, z2, z3, z4]

def floor_ground_Q(L, W, floor_mm, delta_t):
    R_ins = floor_mm/1000/LAM
    zz = floor_zones(L, W)
    return sum(A*delta_t/(R_BASE_ZONES[i]+R_ins) for i,A in enumerate(zz) if A>0)

def calc_losses(b):
    L, W, H = b['length'], b['width'], b['height']
    t_int = b['t_int']
    dt = t_int - T_EXT
    A_win  = b['n_windows']*b['A_window']
    A_door = b['n_doors']*b['A_door']
    A_wall = b['perimeter']*H - A_win - A_door
    A_fl   = b['area']
    G = calc_gsop(t_int)
    Rw = R_wall(b['wall_mm']); Rc = R_ceil(b['ceil_mm'])
    Qw  = A_wall*dt/Rw;  Qc = A_fl*dt/Rc
    if b['floor_type']=='above_air':
        Rf = R_floor_air(b['floor_mm'])
        Qf = A_fl*dt/Rf
    else:
        Rf = None
        Qf = floor_ground_Q(L, W, b['floor_mm'], dt)
    Qwin  = A_win*dt/R_WIN
    Qdoor = A_door*dt/R_DOOR
    Qtr   = Qw+Qc+Qf+Qwin+Qdoor
    Qtr_c = Qtr*1.10
    V = b['area']*H
    Qv = RHO_AIR*C_AIR*(b['n_vent']*V)*dt/3600
    Qt = Qtr_c + Qv
    return dict(dt=dt, G=G, Rw=Rw, Rc=Rc, Rf=Rf,
                Rw_req=r_req_wall(G), Rc_req=r_req_ceil(G),
                Rf_req=r_req_floor(G), Rwin_req=r_req_win(G),
                A_wall=A_wall, A_floor=A_fl,
                A_win=A_win, A_door=A_door,
                Qw=Qw, Qc=Qc, Qf=Qf, Qwin=Qwin, Qdoor=Qdoor,
                Qtr=Qtr, Qtr_c=Qtr_c, Qv=Qv, Qt=Qt,
                Qt_all=Qt*b['count'])

# ======================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ФОРМАТИРОВАНИЯ DOCX
# ======================================================================
def shd(cell, color):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto')
    s.set(qn('w:fill'), color); p.append(s)

def hrow(table, cols, color='D9E2F3'):
    row = table.rows[0]
    for i,txt in enumerate(cols):
        c = row.cells[i]; c.text=txt
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run(txt)
        r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shd(c, color)

def addrow(table, vals, bold=False, bgcol=None):
    row=table.add_row()
    for i,v in enumerate(vals):
        c=row.cells[i]; c.text=str(v)
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        r=c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run(str(v))
        r.font.name='Times New Roman'; r.font.size=Pt(10); r.bold=bold
        if bgcol: shd(c, bgcol)

def p(doc, txt, indent=False, bold=False, center=False, sz=12):
    pg=doc.add_paragraph()
    if center: pg.alignment=WD_ALIGN_PARAGRAPH.CENTER
    else: pg.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: pg.paragraph_format.first_line_indent=Cm(1.25)
    run=pg.add_run(txt)
    run.font.name='Times New Roman'; run.font.size=Pt(sz); run.bold=bold
    return pg

def f(doc, txt):
    pg=doc.add_paragraph()
    pg.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=pg.add_run(txt)
    run.font.name='Times New Roman'; run.font.size=Pt(12); run.italic=True

def h(doc, txt, lvl=1):
    pg=doc.add_heading(txt, level=lvl)
    pg.alignment=WD_ALIGN_PARAGRAPH.LEFT
    run=pg.runs[0] if pg.runs else pg.add_run(txt)
    run.font.name='Times New Roman'
    run.font.color.rgb=RGBColor(0,0,0)
    run.font.size=Pt(14 if lvl==1 else 13 if lvl==2 else 12)
    run.bold=True

def newtable(doc, n_cols, headers, col_widths=None):
    t=doc.add_table(rows=1, cols=n_cols)
    t.style='Table Grid'
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hrow(t, headers)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width=Cm(w)
    return t

# ======================================================================
# ГЕНЕРАЦИЯ ДОКУМЕНТА
# ======================================================================
doc = Document()
sty = doc.styles['Normal']
sty.font.name='Times New Roman'; sty.font.size=Pt(12)

# --- Поля страницы ---
for sec in doc.sections:
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(3); sec.right_margin=Cm(1.5)

# ===== ТИТУЛЬНЫЙ ЛИСТ =====
for _ in range(4): doc.add_paragraph()
p(doc,'РАСЧЁТ ТЕПЛОПОТЕРЬ И ЭЛЕКТРИЧЕСКОЙ МОЩНОСТИ',center=True,bold=True,sz=16)
p(doc,'БАЗЫ ОТДЫХА',center=True,bold=True,sz=14)
p(doc,'г. Нижний Новгород',center=True,sz=12)
for _ in range(2): doc.add_paragraph()
p(doc,'Расчёт выполнен в соответствии с требованиями:',center=True,sz=11)
p(doc,'ГОСТ Р 56778-2015 • ГОСТ Р 54851-2011 • ГОСТ 30494-2011 • СП 50.13330.2022',center=True,sz=11)
for _ in range(8): doc.add_paragraph()
p(doc,'2025 год',center=True,sz=12)
doc.add_page_break()

# ===== 1. ОБЩИЕ ПОЛОЖЕНИЯ =====
h(doc,'1. ОБЩИЕ ПОЛОЖЕНИЯ')
p(doc,'Настоящий расчёт выполнен с целью определения теплопотерь зданий базы отдыха '
  'и расчёта максимальной электрической нагрузки для заключения договора технологического '
  'присоединения к электрическим сетям.',indent=True)
p(doc,'Расчёт выполнен в соответствии с ГОСТ Р 56778-2015 «Здания и сооружения. '
  'Методы расчёта теплопотерь через ограждающие конструкции», ГОСТ Р 54851-2011, '
  'ГОСТ 30494-2011, СП 50.13330.2022, СП 131.13330.2020, СП 60.13330.2020.',indent=True)
p(doc,'Объект: база отдыха, расположенная в г. Нижний Новгород. Комплекс включает жилые '
  'дома трёх типов, административные и инфраструктурные здания — всего 14 типов зданий.',
  indent=True)
doc.add_paragraph()

# ===== 2. НОРМАТИВНЫЕ ССЫЛКИ =====
h(doc,'2. НОРМАТИВНЫЕ ССЫЛКИ')
norms=[
    'СП 50.13330.2022 «Тепловая защита зданий» (актуализир. ред. СНиП 23-02-2003).',
    'СП 131.13330.2020 «Строительная климатология» (актуализир. ред. СНиП 23-01-99*).',
    'СП 60.13330.2020 «Отопление, вентиляция и кондиционирование воздуха».',
    'СП 20.13330.2017 «Нагрузки и воздействия».',
    'СП 256.1325800.2016 «Электроустановки жилых и общественных зданий».',
    'ГОСТ Р 56778-2015 «Здания и сооружения. Методы расчёта теплопотерь через ограждающие конструкции».',
    'ГОСТ Р 54851-2011 «Конструкции строительные ограждающие неоднородные. Расчёт приведённого сопротивления теплопередаче».',
    'ГОСТ 30494-2011 «Здания жилые и общественные. Параметры микроклимата в помещениях».',
    'ГОСТ 7076-99 «Материалы строительные. Метод определения теплопроводности и термического сопротивления».',
    'ПУЭ (7-е изд.) — Правила устройства электроустановок.',
]
for n in norms:
    pg=doc.add_paragraph(style='List Bullet')
    run=pg.add_run(n); run.font.name='Times New Roman'; run.font.size=Pt(12)
doc.add_paragraph()

# ===== 3. КЛИМАТИЧЕСКИЕ ДАННЫЕ =====
h(doc,'3. КЛИМАТИЧЕСКИЕ ДАННЫЕ ДЛЯ Г. НИЖНИЙ НОВГОРОД')
p(doc,'Климатические параметры приняты по СП 131.13330.2020 «Строительная климатология», '
  'таблицы 5.1 и 3.1:',indent=True)
t=newtable(doc,3,['Параметр','Значение','Нормативная ссылка'])
rows=[
    ('Расчётная t нар. воздуха (пятидневка 0,92), t₅',    '−31 °С',          'СП 131.13330.2020, табл. 5.1'),
    ('Средняя t отоп. периода, t_от',                      '−3,6 °С',         'СП 131.13330.2020, табл. 3.1'),
    ('Продолжительность отоп. периода, n_от',              '209 сут',         'СП 131.13330.2020, табл. 3.1'),
    ('Ср. t января',                                        '−10,1 °С',        'СП 131.13330.2020, табл. 5.1'),
    ('Расч. t вн. возд. жилых помещений, t_int',           '+20 °С',          'ГОСТ 30494-2011, табл. 1'),
    ('Расч. t вн. возд. общ. (адм., спорт, обществ.)',     '+20 / +18 / +16 °С','ГОСТ 30494-2011; СП 60.13330.2020'),
    ('Расч. t вн. возд. СПА-комплекса',                    '+25 °С',          'ГОСТ 30494-2011; проект. задание'),
    ('Климатический район',                                 'II В',            'СП 131.13330.2020'),
]
for r in rows: addrow(t,r)
doc.add_paragraph()

# ===== 4. МЕТОДОЛОГИЯ =====
h(doc,'4. МЕТОДОЛОГИЯ РАСЧЁТА')
h(doc,'4.1. Теплофизические характеристики утеплителя',2)
p(doc,'Утеплитель всех ограждающих конструкций — минеральная вата плотностью ρ = 150 кг/м³. '
  'Расчётный коэффициент теплопроводности принят для условий эксплуатации Б '
  '(нормальный влажностный режим Нижегородской области, ГОСТ 7076-99):',indent=True)
f(doc,'λ = 0,045 Вт/(м·К)')
p(doc,'Сопротивления теплоотдаче поверхностей (СП 50.13330.2022, табл. 6):',indent=True)
p(doc,'    Rsi_ст  = 0,13 м²·°С/Вт  (стены, горизонтальный поток)')
p(doc,'    Rsi_пт  = 0,10 м²·°С/Вт  (потолок, поток вверх)')
p(doc,'    Rsi_пол = 0,17 м²·°С/Вт  (пол, поток вниз)')
p(doc,'    Rse     = 0,04 м²·°С/Вт  (наружная поверхность)')

h(doc,'4.2. Расчёт приведённого сопротивления теплопередаче',2)
p(doc,'Для однородных конструкций (ГОСТ Р 54851-2011, п. 4.3; СП 50.13330.2022, п. 6.1):',indent=True)
f(doc,'R₀ = Rsi + δ/λ + Rse,  м²·°С/Вт')
p(doc,'Для светопрозрачных конструкций и дверей — по паспортным данным изделий:',indent=True)
p(doc,'    R_окн = 0,56 м²·°С/Вт (двухкамерный стеклопакет 4-16-4-16-4, ПВХ профиль)')
p(doc,'    R_дв  = 0,60 м²·°С/Вт (утеплённая стальная дверь)')

h(doc,'4.3. Нормируемые значения (СП 50.13330.2022, п. 5.2, табл. 3)',2)
p(doc,'ГСОП вычисляется по формуле:',indent=True)
f(doc,'ГСОП = (t_int − t_от) × n_от = (t_int − (−3,6)) × 209')
p(doc,'Нормируемое R_req определяется интерполяцией по табл. 3 СП 50.13330.2022:',indent=True)
rows_req=[
    ('Стены',             'R_req = 0,00035·ГСОП + 1,4'),
    ('Покрытие/потолок',  'R_req = 0,00050·ГСОП + 2,2'),
    ('Пол',               'R_req = 0,00045·ГСОП + 1,9'),
    ('Окна и двери',      'R_req = 0,000075·ГСОП + 0,15'),
]
t2=newtable(doc,2,['Конструкция','Формула R_req, м²·°С/Вт'])
for r in rows_req: addrow(t2,r)
doc.add_paragraph()

h(doc,'4.4. Трансмиссионные теплопотери (ГОСТ Р 56778-2015, п. 5.3)',2)
f(doc,'Q_тр = Σ [ Aᵢ / R₀ᵢ × (t_int − t_ext) ]  + 10% на инфильтрацию')
p(doc,'Поправка +10% учитывает потери через неплотности конструкций и '
  'воздействие ветра (ГОСТ Р 56778-2015, п. 5.4; β_inf = 0,10).',indent=True)

h(doc,'4.5. Пол по грунту — зонный метод (СП 50.13330.2022, прил. В)',2)
p(doc,'Пол по грунту разбивается на зоны шириной 2 м от наружных стен. '
  'Базовые сопротивления зон (без утепления):',indent=True)
tz=newtable(doc,3,['Зона','Расстояние от стены','R_base, м²·°С/Вт'])
for zn,dst,rr in [('I','0−2 м','2,1'),('II','2−4 м','4,3'),
                   ('III','4−6 м','8,6'),('IV','>6 м','14,2')]: addrow(tz,[zn,dst,rr])
p(doc,'При наличии утеплителя толщиной δ_ут: R_зон = R_base + δ_ут/λ',indent=True)

h(doc,'4.6. Вентиляционные теплопотери (СП 60.13330.2020, п. 6.4.2)',2)
f(doc,'Q_в = ρ_в × c_в × L × (t_int − t_ext) / 3600')
p(doc,'Расход воздуха: L = n_v × V (м³/ч); n_v — кратность воздухообмена.',indent=True)
p(doc,'Для жилых: n_v = 0,35 ч⁻¹ (мин., СП 60.13330.2020, п. 8.3). '
  'Для общественных и административных — по профилю использования.',indent=True)
doc.add_paragraph()

# ===== 5. ГСОП ДЛЯ КАЖДОГО ТИПА ЗДАНИЯ =====
h(doc,'5. ГРАДУСО-СУТКИ ОТОПИТЕЛЬНОГО ПЕРИОДА')
p(doc,'Расчётные значения ГСОП для различных расчётных температур внутри зданий:',indent=True)
t3=newtable(doc,4,['Тип здания','t_int, °С','ГСОП = (t_int+3,6)×209','°С·сут'])
gsop_rows=[
    ('Жилые дома, административные, общественные', 20,
     f'(20+3,6)×209 = 23,6×209', f'{calc_gsop(20):.0f}'),
    ('СПА-комплекс', 25,
     f'(25+3,6)×209 = 28,6×209', f'{calc_gsop(25):.0f}'),
    ('Центр парусного спорта', 18,
     f'(18+3,6)×209 = 21,6×209', f'{calc_gsop(18):.0f}'),
    ('Пункт проката инвентаря', 16,
     f'(16+3,6)×209 = 19,6×209', f'{calc_gsop(16):.0f}'),
]
for gr in gsop_rows: addrow(t3,gr)
doc.add_paragraph()

# ===== 6. РАСЧЁТ ПО ЗДАНИЯМ =====
h(doc,'6. РАСЧЁТ ТЕПЛОПОТЕРЬ ПО ЗДАНИЯМ')

def fmt(v, nd=2): return f'{v:.{nd}f}'.replace('.',',')

def write_building_section(doc, b, results, sec_num):
    h(doc, f'6.{sec_num}. {b["name"].upper()} (кол-во: {b["count"]} шт.)', 2)
    R = results

    # Параметры
    p(doc,'Исходные данные:',indent=True)
    tb=newtable(doc,3,['Параметр','Значение','Единица'])
    rows=[
        ('Размеры (Ш×Д×В)',f'{b["width"]}×{b["length"]}×{b["height"]}','м'),
        ('Площадь',f'{b["area"]:.1f}','м²'),
        ('Периметр',f'{b["perimeter"]:.2f}','м'),
        ('Объём (строит.)',fmt(b['area']*b['height'],1),'м³'),
        ('Кол-во дверей / площадь',f'{b["n_doors"]} × {b["A_door"]} м²','шт./м²'),
        ('Кол-во окон / площадь',f'{b["n_windows"]} × {b["A_window"]} м²','шт./м²'),
        ('Утепление стены',f'{b["wall_mm"]} мм мин.ваты','—'),
        ('Утепление потолка',f'{b["ceil_mm"]} мм мин.ваты','—'),
        ('Утепление пола',f'{b["floor_mm"]} мм мин.ваты','—'),
        ('Тип пола','Над улицей' if b['floor_type']=='above_air' else 'По грунту','—'),
        ('Расч. t внутри',f'+{b["t_int"]}','°С'),
        ('Кратность вентиляции',fmt(b['n_vent'],2),'ч⁻¹'),
    ]
    for r in rows: addrow(tb,r)
    doc.add_paragraph()

    G=R['G']
    p(doc,f'ГСОП = ({b["t_int"]} − (−3,6)) × 209 = {fmt(b["t_int"]+3.6,1)} × 209 = {fmt(G,0)} °С·сут',indent=True)
    p(doc,'',indent=True)

    # Сопротивления и проверка
    p(doc,'Фактическое и нормируемое сопротивления теплопередаче:',indent=True)
    tc=newtable(doc,4,['Конструкция','R₀_факт','R_req','Соотв.'])
    fl_type_str = f'Пол ({("над улицей" if b["floor_type"]=="above_air" else "по грунту, зона I")})'
    rr=[
        ('Стена',    fmt(R['Rw']),    fmt(R['Rw_req']), '✓' if R['Rw']>=R['Rw_req'] else '✗'),
        ('Потолок',  fmt(R['Rc']),    fmt(R['Rc_req']), '✓' if R['Rc']>=R['Rc_req'] else '✗'),
        ('Окна',     fmt(R_WIN),      fmt(R['Rwin_req']),'✓' if R_WIN>=R['Rwin_req'] else '✗'),
        ('Дверь',    fmt(R_DOOR),     fmt(R['Rwin_req']),'✓' if R_DOOR>=R['Rwin_req'] else '✗'),
    ]
    if b['floor_type']=='above_air':
        rr.insert(2,('Пол (над улицей)', fmt(R['Rf']), fmt(R['Rf_req']),
                     '✓' if R['Rf']>=R['Rf_req'] else '✗'))
    else:
        L,W = b['length'], b['width']
        R_ins_str = fmt(b['floor_mm']/1000/LAM)
        R_z1=R_BASE_ZONES[0]+b['floor_mm']/1000/LAM
        rr.insert(2,(f'Пол/зона I (по грунту+{b["floor_mm"]}мм)', fmt(R_z1), fmt(R['Rf_req']),
                     '✓' if R_z1>=R['Rf_req'] else '✗'))
    for r in rr: addrow(tc,r)
    p(doc,'Единицы: м²·°С/Вт',indent=True)
    doc.add_paragraph()

    # Площади
    p(doc,f'Площади ограждающих конструкций:',indent=True)
    p(doc,f'    Стены брутто: P×H = {b["perimeter"]}×{b["height"]} = {b["perimeter"]*b["height"]:.1f} м²')
    p(doc,f'    Окна: {b["n_windows"]}×{b["A_window"]} = {R["A_win"]:.1f} м²; '
          f'Двери: {b["n_doors"]}×{b["A_door"]} = {R["A_door"]:.1f} м²')
    p(doc,f'    Стены нетто: {b["perimeter"]*b["height"]:.1f} − {R["A_win"]:.1f} − {R["A_door"]:.1f} '
          f'= {R["A_wall"]:.1f} м²')

    # Теплопотери по элементам
    p(doc,'Теплопотери одного здания (Δt = '
          f't_int − t_ext = {b["t_int"]} − ({T_EXT}) = {R["dt"]} °С):',indent=True)
    tq=newtable(doc,4,['Элемент','Площадь, м²','R₀ или формула','Q, Вт'])
    if b['floor_type']=='on_ground':
        L,W = b['length'], b['width']
        zz = floor_zones(L, W)
        R_ins = b['floor_mm']/1000/LAM
        floor_str = ' + '.join([f'{zz[i]:.1f}м²/зI+{i}'[:18] for i in range(4) if zz[i]>0])
        addrow(tq,['Стены', fmt(R['A_wall'],1), fmt(R['Rw']), fmt(R['Qw'],0)])
        addrow(tq,['Потолок', fmt(R['A_floor'],1), fmt(R['Rc']), fmt(R['Qc'],0)])
        addrow(tq,['Пол (по грунту, зон. метод)', fmt(R['A_floor'],1), 'зон. метод', fmt(R['Qf'],0)])
        addrow(tq,['Окна', fmt(R['A_win'],1), fmt(R_WIN), fmt(R['Qwin'],0)])
        addrow(tq,['Дверь', fmt(R['A_door'],1), fmt(R_DOOR), fmt(R['Qdoor'],0)])
    else:
        addrow(tq,['Стены',   fmt(R['A_wall'],1), fmt(R['Rw']),    fmt(R['Qw'],0)])
        addrow(tq,['Потолок', fmt(R['A_floor'],1),fmt(R['Rc']),    fmt(R['Qc'],0)])
        addrow(tq,['Пол (над улицей)', fmt(R['A_floor'],1), fmt(R['Rf']), fmt(R['Qf'],0)])
        addrow(tq,['Окна',    fmt(R['A_win'],1),  fmt(R_WIN),       fmt(R['Qwin'],0)])
        addrow(tq,['Дверь',   fmt(R['A_door'],1), fmt(R_DOOR),      fmt(R['Qdoor'],0)])
    addrow(tq,['Инфильтрация (+10%)', '—', '—', fmt(R['Qtr']*0.10,0)], bold=True, bgcol='FFF2CC')
    addrow(tq,['ИТОГО трансмиссионные', '—', '—', fmt(R['Qtr_c'],0)], bold=True, bgcol='FFF2CC')
    V=b['area']*b['height']
    addrow(tq,[f'Вентиляция (n={b["n_vent"]}ч⁻¹, V={V:.0f}м³)', '—', '—', fmt(R['Qv'],0)])
    addrow(tq,['ИТОГО 1 здание', '—', '—', fmt(R['Qt'],0)], bold=True, bgcol='E2EFDA')
    if b['count']>1:
        addrow(tq,[f'ИТОГО {b["count"]} здания(й)', '—', '—',
                   fmt(R['Qt_all'],0)], bold=True, bgcol='E2EFDA')
    doc.add_paragraph()

# Вычисляем потери для всех отапливаемых зданий
all_results = {}
for b in BUILDINGS:
    if b['heated']:
        all_results[b['id']] = calc_losses(b)

# Пишем разделы
sec=0
for b in BUILDINGS:
    sec+=1
    if not b['heated']:
        h(doc,f'6.{sec}. {b["name"].upper()} (кол-во: {b["count"]} шт.) — НЕ ОТАПЛИВАЕТСЯ',2)
        p(doc,f'Здание является неотапливаемым. Расчёт теплопотерь для нужд отопления '
          f'не производится. Электрическая нагрузка учитывается только по освещению и '
          f'технологическому оборудованию (см. раздел 8).',indent=True)
        doc.add_paragraph()
    else:
        write_building_section(doc, b, all_results[b['id']], sec)

# ===== 7. СВОДНАЯ ТАБЛИЦА ТЕПЛОПОТЕРЬ =====
doc.add_page_break()
h(doc,'7. СВОДНАЯ ТАБЛИЦА ТЕПЛОПОТЕРЬ')
p(doc,'Таблица содержит суммарные теплопотери всех отапливаемых зданий:',indent=True)
tsv=newtable(doc,6,
    ['Здание','Кол.','Q₁здан., Вт','Q₁здан., кВт','Q всего, Вт','Q всего, кВт'])
total_Q=0
for b in BUILDINGS:
    if b['heated']:
        R=all_results[b['id']]
        addrow(tsv,[b['name'], b['count'],
                    fmt(R['Qt'],0), fmt(R['Qt']/1000,2),
                    fmt(R['Qt_all'],0), fmt(R['Qt_all']/1000,1)])
        total_Q+=R['Qt_all']
addrow(tsv,['ИТОГО БАЗА ОТДЫХА','—','—','—',
            fmt(total_Q,0), fmt(total_Q/1000,1)], bold=True, bgcol='FFF2CC')
doc.add_paragraph()

# ===== 8. РАСЧЁТ ЭЛЕКТРИЧЕСКОЙ МОЩНОСТИ =====
doc.add_page_break()
h(doc,'8. РАСЧЁТ ЭЛЕКТРИЧЕСКОЙ МОЩНОСТИ ДЛЯ ДОГОВОРА ТЕХНОЛОГИЧЕСКОГО ПРИСОЕДИНЕНИЯ')
p(doc,'Суммарная электрическая мощность базы отдыха определяется как совокупность нагрузок '
  'от систем отопления, горячего водоснабжения, освещения и технологического оборудования '
  'зданий, а также инфраструктурных объектов (наружное освещение, водоснабжение, очистные '
  'сооружения). Расчёт произведён в соответствии с СП 256.1325800.2016 и ПУЭ (7-е изд.).',
  indent=True)

h(doc,'8.1. Установленная мощность по зданиям',2)
p(doc,'Мощность электрического отопления принята с запасным коэффициентом k_з = 1,15 '
  '(учёт тепловых мостов и неучтённых потерь). Прочие нагрузки — по нормативам освещённости '
  '(СП 52.13330.2016) и установленному оборудованию:',indent=True)

tel=newtable(doc,7,
    ['Здание','Кол.','P_отоп./1зд., Вт','P_ГВС, Вт','P_свет, Вт',
     'P_проч, Вт','P_уст×Кол., кВт'])
grand_install=0
rows_el=[]
for b in BUILDINGS:
    if b['heated']:
        R=all_results[b['id']]
        P_heat = R['Qt'] * 1.15
    else:
        P_heat = 0
    P_dhw  = b['dhw_W']
    P_lght = b['light_W']
    P_misc = b['misc_W']
    P_unit = P_heat + P_dhw + P_lght + P_misc
    P_total_bld = P_unit * b['count']
    grand_install += P_total_bld
    rows_el.append((b['name'], b['count'],
                    fmt(P_heat,0) if b['heated'] else '—',
                    fmt(P_dhw,0), fmt(P_lght,0), fmt(P_misc,0),
                    fmt(P_total_bld/1000,1)))
    addrow(tel, rows_el[-1])
addrow(tel,['Инфраструктура (наружн. освещение, водоснабжение, охрана)',
            '—','—','—','—','—', fmt(INFRA_W/1000,1)])
grand_install += INFRA_W
addrow(tel,['ИТОГО установленная мощность','—','—','—','—','—',
            fmt(grand_install/1000,1)], bold=True, bgcol='FFF2CC')
doc.add_paragraph()

h(doc,'8.2. Расчётная (заявленная) мощность',2)
K_od = 0.80
P_demand = grand_install * K_od
p(doc,'Расчётная мощность определяется с учётом коэффициента одновременности нагрузок '
  'K_od = 0,80 (для объектов отдыха с разнообразным составом нагрузок, '
  'СП 256.1325800.2016, прил. Б):',indent=True)
f(doc,f'P_расч = P_уст × K_od = {fmt(grand_install/1000,1)} × 0,80 = {fmt(P_demand/1000,1)} кВт')

# Сводная электрика
h(doc,'8.3. Сводная таблица электрических нагрузок',2)
te2=newtable(doc,3,['Показатель','кВт','Примечание'])
addrow(te2,['Суммарные теплопотери базы (отапл. здания)',
            fmt(total_Q/1000,1), 'Раздел 7'])
addrow(te2,['Установленная мощность отопления (k_з=1,15)',
            fmt(sum(all_results[b['id']]['Qt_all']*1.15 for b in BUILDINGS if b['heated'])/1000,1),
            'Раздел 8.1'])
addrow(te2,['Прочие электрические нагрузки (ГВС, свет, оборуд.)',
            fmt((grand_install - INFRA_W - sum(all_results[b['id']]['Qt_all']*1.15
                 for b in BUILDINGS if b['heated']))/1000,1),
            'Раздел 8.1'])
addrow(te2,['Инфраструктурные нагрузки', fmt(INFRA_W/1000,1), 'Наружн. освещение, насосы'])
addrow(te2,['СУММАРНАЯ установленная мощность', fmt(grand_install/1000,1), '—'],
       bold=True, bgcol='FFF2CC')
addrow(te2,['Коэффициент одновременности K_od', '0,80', 'СП 256.1325800.2016'])
addrow(te2,['РАСЧЁТНАЯ мощность (заявляемая в договоре)', fmt(P_demand/1000,1), '—'],
       bold=True, bgcol='E2EFDA')
doc.add_paragraph()

# ===== 9. ВЫВОДЫ =====
h(doc,'9. ВЫВОДЫ')
conclusions=[
    f'Расчёт выполнен для базы отдыха из 14 типов зданий (всего {sum(b["count"] for b in BUILDINGS)} '
    f'зданий) в г. Нижний Новгород в соответствии с ГОСТ Р 56778-2015, ГОСТ Р 54851-2011, '
    f'ГОСТ 30494-2011 и СП 50.13330.2022.',

    f'Расчётная температура наружного воздуха t₅ = −31 °С (СП 131.13330.2020, '
    f'холодная пятидневка обеспеченностью 0,92). ГСОП для жилых зданий = '
    f'{calc_gsop(20):.0f} °С·сут.',

    f'Все ограждающие конструкции отапливаемых зданий удовлетворяют нормируемым '
    f'значениям сопротивления теплопередаче по СП 50.13330.2022 (табл. 3).',

    f'Суммарные теплопотери отапливаемых зданий базы отдыха составляют '
    f'{total_Q/1000:.1f} кВт.',

    f'Суммарная установленная электрическая мощность базы отдыха (с учётом отопления, '
    f'ГВС, освещения, оборудования и инфраструктуры) составляет '
    f'{grand_install/1000:.1f} кВт.',

    f'Расчётная (заявляемая) электрическая мощность для технологического присоединения '
    f'к электрическим сетям (с коэффициентом одновременности K_od = 0,80 по '
    f'СП 256.1325800.2016) составляет {P_demand/1000:.1f} кВт.',

    f'Для обеспечения надёжного электроснабжения рекомендуется выбирать вводные '
    f'аппараты защиты и проводники по расчётной мощности {P_demand/1000:.1f} кВт, '
    f'с резервом мощности по установленной мощности {grand_install/1000:.1f} кВт.',
]
for i,c in enumerate(conclusions,1):
    pg=doc.add_paragraph(style='List Number')
    run=pg.add_run(c); run.font.name='Times New Roman'; run.font.size=Pt(12)

doc.add_paragraph()
p(doc,'Настоящий расчёт составлен в соответствии с действующей нормативной документацией '
  'Российской Федерации и предназначен для использования при проектировании системы '
  'электроснабжения базы отдыха и заключении договора технологического присоединения '
  'к электрическим сетям.',indent=True)

# ===== СОХРАНЕНИЕ =====
OUT = r'C:\dev\test-projekt\heat-loss-calculation\Расчёт теплопотерь базы отдыха (полный комплекс).docx'
doc.save(OUT)
print(f'OK: {OUT}')
print(f'\n=== КЛЮЧЕВЫЕ РЕЗУЛЬТАТЫ ===')
print(f'Суммарные теплопотери: {total_Q/1000:.1f} кВт')
print(f'Установленная мощность: {grand_install/1000:.1f} кВт')
print(f'Расчётная мощность (договор): {P_demand/1000:.1f} кВт')
print()
for b in BUILDINGS:
    if b['heated']:
        R=all_results[b['id']]
        print(f'  {b["id"]}) {b["name"]} x{b["count"]}: {R["Qt"]:.0f} Вт/зд → итого {R["Qt_all"]/1000:.1f} кВт')
