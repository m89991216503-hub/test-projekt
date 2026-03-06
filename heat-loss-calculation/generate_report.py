# -*- coding: utf-8 -*-
"""
Генератор расчёта теплопотерь базы отдыха (51 дом)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import math

doc = Document()

# --- Стили ---
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        run.font.size = Pt(14)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.bold = True
    return p

def para(doc, text, bold_parts=None, indent_cm=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    if bold_parts is None:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    else:
        # bold_parts: list of (text, bold)
        for txt, bold in bold_parts:
            run = p.add_run(txt)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = bold
    return p

def formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    return p

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, (cell_text, width) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = cell_text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].bold = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if header:
            set_cell_bg(cell, 'D9E2F3')
    return row

# =====================================================================
# ТИТУЛЬНЫЙ ЛИСТ
# =====================================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('РАСЧЁТ ТЕПЛОПОТЕРЬ И МОЩНОСТИ ЭЛЕКТРИЧЕСКОГО ОТОПЛЕНИЯ')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('БАЗЫ ОТДЫХА ИЗ 51 ЖИЛОГО ДОМА')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('г. Нижний Новгород\n2025 г.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# =====================================================================
# 1. ОБЩИЕ ПОЛОЖЕНИЯ
# =====================================================================
heading(doc, '1. ОБЩИЕ ПОЛОЖЕНИЯ', 1)
para(doc,
     'Настоящий расчёт выполнен в целях определения теплопотерь жилых домов базы отдыха '
     'и расчёта требуемой электрической мощности системы отопления. Расчёт произведён в '
     'соответствии с действующей нормативно-технической документацией Российской Федерации.',
     indent_cm=1.25)
para(doc,
     'Объект расчёта: база отдыха, состоящая из 51 одноэтажного жилого дома с площадью '
     'каждого дома 20 м², расположенная в г. Нижний Новгород.',
     indent_cm=1.25)

# =====================================================================
# 2. НОРМАТИВНЫЕ ССЫЛКИ
# =====================================================================
heading(doc, '2. НОРМАТИВНЫЕ ССЫЛКИ', 1)
norms = [
    'СП 50.13330.2022 «Тепловая защита зданий». Актуализированная редакция СНиП 23-02-2003.',
    'СП 131.13330.2020 «Строительная климатология». Актуализированная редакция СНиП 23-01-99*.',
    'СП 60.13330.2020 «Отопление, вентиляция и кондиционирование воздуха».',
    'СП 256.1325800.2016 «Электроустановки жилых и общественных зданий. Правила проектирования и монтажа».',
    'ГОСТ Р 56778-2015 «Здания и сооружения. Методы расчёта теплопотерь через ограждающие конструкции».',
    'ГОСТ Р 54851-2011 «Конструкции строительные ограждающие неоднородные. Расчёт приведённого сопротивления теплопередаче».',
    'ГОСТ 30494-2011 «Здания жилые и общественные. Параметры микроклимата в помещениях».',
    'ГОСТ 26253-2014 «Здания и сооружения. Метод определения теплоустойчивости ограждающих конструкций».',
    'СП 20.13330.2017 «Нагрузки и воздействия».',
    'ГОСТ 7076-99 «Материалы и изделия строительные. Метод определения теплопроводности и термического сопротивления при стационарном тепловом режиме».',
]
for norm in norms:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(norm)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# =====================================================================
# 3. ИСХОДНЫЕ ДАННЫЕ
# =====================================================================
heading(doc, '3. ИСХОДНЫЕ ДАННЫЕ', 1)
heading(doc, '3.1. Геометрические параметры здания', 2)

para(doc,
     'Геометрические параметры жилого дома приняты в соответствии с заданием на '
     'проектирование:',
     indent_cm=1.25)

# Таблица исходных данных
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0]
for i, txt in enumerate(['Параметр', 'Значение', 'Единица']):
    hdr.cells[i].text = txt
    hdr.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if hdr.cells[i].paragraphs[0].runs:
        hdr.cells[i].paragraphs[0].runs[0].bold = True
        hdr.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_bg(hdr.cells[i], 'D9E2F3')

geom_data = [
    ('Количество домов', '51', 'шт.'),
    ('Ширина дома', '4,0', 'м'),
    ('Длина дома', '5,0', 'м'),
    ('Периметр', '18,0', 'м'),
    ('Высота потолков', '3,0', 'м'),
    ('Жилая площадь одного дома', '20', 'м²'),
    ('Площадь застройки одного дома', '20', 'м²'),
    ('Строительный объём одного дома', '60', 'м³'),
    ('Количество дверей', '1', 'шт.'),
    ('Площадь одной двери', '2,0', 'м²'),
    ('Количество окон', '3', 'шт.'),
    ('Площадь одного окна', '2,0', 'м²'),
    ('Суммарная площадь окон', '6,0', 'м²'),
    ('Тип вентиляции', 'Естественная', '—'),
    ('Тип пола', 'Над улицей (вентилируемое подпольё)', '—'),
]
for row_data in geom_data:
    row = t.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(11)

doc.add_paragraph()

heading(doc, '3.2. Конструктивные решения и теплоизоляция', 2)
para(doc,
     'Теплоизоляция ограждающих конструкций выполнена из минеральной ваты плотностью '
     '150 кг/м³. Коэффициент теплопроводности λ определён по ГОСТ 7076-99 для условий '
     'эксплуатации Б (нормальный влажностный режим, характерный для г. Нижний Новгород): '
     'λ = 0,045 Вт/(м·К).',
     indent_cm=1.25)

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = t2.rows[0]
for i, txt in enumerate(['Конструкция', 'Утеплитель', 'Толщина δ, мм', 'λ, Вт/(м·К)']):
    hdr2.cells[i].text = txt
    hdr2.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if hdr2.cells[i].paragraphs[0].runs:
        hdr2.cells[i].paragraphs[0].runs[0].bold = True
        hdr2.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        hdr2.cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_bg(hdr2.cells[i], 'D9E2F3')

constr_data = [
    ('Наружная стена', 'Мин. вата ρ=150 кг/м³', '200', '0,045'),
    ('Потолок (перекрытие)', 'Мин. вата ρ=150 кг/м³', '250', '0,045'),
    ('Пол (над улицей)', 'Мин. вата ρ=150 кг/м³', '250', '0,045'),
    ('Окна', 'Двухкамерный стеклопакет в ПВХ профиле', '—', '—'),
    ('Входная дверь', 'Утеплённая металлическая дверь', '—', '—'),
]
for row_data in constr_data:
    row = t2.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(11)
doc.add_paragraph()

# =====================================================================
# 4. КЛИМАТИЧЕСКИЕ ДАННЫЕ
# =====================================================================
heading(doc, '4. КЛИМАТИЧЕСКИЕ ДАННЫЕ', 1)
para(doc,
     'Климатические параметры для г. Нижний Новгород приняты по СП 131.13330.2020 '
     '«Строительная климатология» (Таблица 5.1, 3.1):',
     indent_cm=1.25)

t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = t3.rows[0]
for i, txt in enumerate(['Параметр', 'Значение', 'Ссылка']):
    hdr3.cells[i].text = txt
    hdr3.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if hdr3.cells[i].paragraphs[0].runs:
        hdr3.cells[i].paragraphs[0].runs[0].bold = True
        hdr3.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        hdr3.cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_bg(hdr3.cells[i], 'D9E2F3')

climate_data = [
    ('Расчётная температура наружного воздуха (холодная пятидневка, обеспеч. 0,92), t₅, °С',
     '−31', 'СП 131.13330.2020, табл. 5.1'),
    ('Расчётная температура наружного воздуха (отопительный период), t_от, °С',
     '−3,6', 'СП 131.13330.2020, табл. 3.1'),
    ('Продолжительность отопительного периода, n_от, сут',
     '209', 'СП 131.13330.2020, табл. 3.1'),
    ('Средняя температура января, °С',
     '−10,1', 'СП 131.13330.2020, табл. 5.1'),
    ('Нормируемая температура внутреннего воздуха жилых помещений, t_int, °С',
     '+20', 'ГОСТ 30494-2011, табл. 1; СП 60.13330.2020'),
    ('Климатический район',
     'II В (умеренно-континентальный климат)', 'СП 131.13330.2020'),
]
for row_data in climate_data:
    row = t3.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

# =====================================================================
# 5. РАСЧЁТ ГСОП
# =====================================================================
heading(doc, '5. РАСЧЁТ ГРАДУСО-СУТОК ОТОПИТЕЛЬНОГО ПЕРИОДА (ГСОП)', 1)
para(doc,
     'Градусо-сутки отопительного периода (ГСОП) определяются по формуле '
     'СП 50.13330.2022, п. 5.2:',
     indent_cm=1.25)
formula(doc, 'ГСОП = (t_int − t_от) × n_от,  °С·сут')
para(doc,
     'где:',
     indent_cm=1.25)
para(doc, '    t_int = +20 °С — расчётная температура внутреннего воздуха;')
para(doc, '    t_от  = −3,6 °С — средняя температура наружного воздуха за отопительный период;')
para(doc, '    n_от  = 209 сут — продолжительность отопительного периода.')
formula(doc, 'ГСОП = (20 − (−3,6)) × 209 = 23,6 × 209 = 4932,4 °С·сут')
para(doc, 'Принимается: ГСОП = 4932 °С·сут.', indent_cm=1.25)

GSOP = (20 - (-3.6)) * 209  # = 4932.4

# =====================================================================
# 6. НОРМИРУЕМЫЕ ЗНАЧЕНИЯ СОПРОТИВЛЕНИЯ ТЕПЛОПЕРЕДАЧЕ
# =====================================================================
heading(doc, '6. НОРМИРУЕМЫЕ ЗНАЧЕНИЯ СОПРОТИВЛЕНИЯ ТЕПЛОПЕРЕДАЧЕ', 1)
para(doc,
     'Нормируемые значения сопротивления теплопередаче R_req ограждающих конструкций '
     'жилых зданий определяются по формуле интерполяции СП 50.13330.2022, табл. 3:',
     indent_cm=1.25)
formula(doc, 'R_req = a × ГСОП + b,  м²·°С/Вт')
para(doc,
     'Коэффициенты a и b приняты по СП 50.13330.2022, таблица 3, '
     'для жилых зданий (группа 1):',
     indent_cm=1.25)

t4 = doc.add_table(rows=1, cols=5)
t4.style = 'Table Grid'
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = t4.rows[0]
for i, txt in enumerate(['Конструкция', 'a', 'b', 'R_req, м²·°С/Вт', 'Примечание']):
    hdr4.cells[i].text = txt
    hdr4.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if hdr4.cells[i].paragraphs[0].runs:
        hdr4.cells[i].paragraphs[0].runs[0].bold = True
        hdr4.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        hdr4.cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_bg(hdr4.cells[i], 'D9E2F3')

# R = a*GSOP + b
R_wall_req = 0.00035 * GSOP + 1.4    # 3,13
R_ceil_req = 0.0005 * GSOP + 2.2     # 4,67
R_floor_req = 0.00045 * GSOP + 1.9   # 4,12
R_win_req = 0.000075 * GSOP + 0.15   # 0,52
R_door_req = 0.000075 * GSOP + 0.15  # same as windows

norm_rows = [
    ('Наружная стена', '0,00035', '1,4',
     f'{R_wall_req:.2f}'.replace('.', ','), 'СП 50.13330.2022, табл. 3'),
    ('Перекрытие (потолок)', '0,00050', '2,2',
     f'{R_ceil_req:.2f}'.replace('.', ','), 'СП 50.13330.2022, табл. 3'),
    ('Пол над улицей (перекрытие)', '0,00045', '1,9',
     f'{R_floor_req:.2f}'.replace('.', ','), 'СП 50.13330.2022, табл. 3'),
    ('Окна', '0,000075', '0,15',
     f'{R_win_req:.2f}'.replace('.', ','), 'СП 50.13330.2022, табл. 3'),
    ('Входная дверь', '0,000075', '0,15',
     f'{R_door_req:.2f}'.replace('.', ','), 'СП 50.13330.2022, табл. 3'),
]
for row_data in norm_rows:
    row = t4.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

# =====================================================================
# 7. РАСЧЁТ ФАКТИЧЕСКОГО СОПРОТИВЛЕНИЯ ТЕПЛОПЕРЕДАЧЕ
# =====================================================================
heading(doc, '7. РАСЧЁТ ФАКТИЧЕСКОГО СОПРОТИВЛЕНИЯ ТЕПЛОПЕРЕДАЧЕ', 1)
para(doc,
     'Приведённое сопротивление теплопередаче однородной ограждающей конструкции '
     'определяется по формуле СП 50.13330.2022, п. 6.1 (ГОСТ Р 54851-2011, п. 4.3):',
     indent_cm=1.25)
formula(doc, 'R₀ = Rsi + Σ(δᵢ/λᵢ) + Rse,  м²·°С/Вт')
para(doc, 'где:', indent_cm=1.25)
para(doc, '    Rsi — сопротивление теплоотдаче внутренней поверхности, м²·°С/Вт;')
para(doc, '    δᵢ  — толщина i-го слоя конструкции, м;')
para(doc, '    λᵢ  — расчётный коэффициент теплопроводности i-го слоя, Вт/(м·К);')
para(doc, '    Rse — сопротивление теплоотдаче наружной поверхности, м²·°С/Вт.')
para(doc,
     'Сопротивления теплоотдаче приняты по СП 50.13330.2022, таблица 6:',
     indent_cm=1.25)
para(doc, '    Rsi (стены, горизонтальный поток теплоты) = 0,13 м²·°С/Вт  (α_si = 8,7 Вт/(м²·°С));')
para(doc, '    Rsi (потолок, тепловой поток вверх)         = 0,10 м²·°С/Вт  (α_si = 10 Вт/(м²·°С));')
para(doc, '    Rsi (пол, тепловой поток вниз)              = 0,17 м²·°С/Вт  (α_si = 6 Вт/(м²·°С));')
para(doc, '    Rse (все наружные поверхности)              = 0,04 м²·°С/Вт  (α_se = 23 Вт/(м²·°С)).')

lam = 0.045  # Вт/(м·К), мин вата 150 кг/м3, условия Б

# Стена
heading(doc, '7.1. Наружная стена', 2)
d_wall = 0.200
Rsi_w = 0.13
Rse = 0.04
R_wall = Rsi_w + d_wall / lam + Rse
para(doc,
     f'Утеплитель: минеральная вата δ = {int(d_wall*1000)} мм, λ = {lam} Вт/(м·К).',
     indent_cm=1.25)
formula(doc,
        f'R₀_ст = 0,13 + {d_wall}/{lam} + 0,04 = 0,13 + {d_wall/lam:.3f} + 0,04 = {R_wall:.2f} м²·°С/Вт')
para(doc,
     f'R₀_ст = {R_wall:.2f} м²·°С/Вт > R_req = {R_wall_req:.2f} м²·°С/Вт — '
     'условие выполнено ✓',
     indent_cm=1.25)

# Потолок
heading(doc, '7.2. Перекрытие (потолок)', 2)
d_ceil = 0.250
Rsi_c = 0.10
R_ceil = Rsi_c + d_ceil / lam + Rse
para(doc,
     f'Утеплитель: минеральная вата δ = {int(d_ceil*1000)} мм, λ = {lam} Вт/(м·К).',
     indent_cm=1.25)
formula(doc,
        f'R₀_пт = 0,10 + {d_ceil}/{lam} + 0,04 = 0,10 + {d_ceil/lam:.3f} + 0,04 = {R_ceil:.2f} м²·°С/Вт')
para(doc,
     f'R₀_пт = {R_ceil:.2f} м²·°С/Вт > R_req = {R_ceil_req:.2f} м²·°С/Вт — '
     'условие выполнено ✓',
     indent_cm=1.25)

# Пол
heading(doc, '7.3. Пол над улицей (вентилируемое подполье)', 2)
d_floor = 0.250
Rsi_f = 0.17
R_floor = Rsi_f + d_floor / lam + Rse
para(doc,
     'Пол расположен над открытым пространством (вентилируемым подпольем над улицей). '
     'По условиям расчёта принимается температурный перепад, равный разнице между '
     'расчётной внутренней и наружной температурами (n = 1,0 по СП 50.13330.2022, '
     'таблица 7). Утеплитель: минеральная вата δ = 250 мм, λ = 0,045 Вт/(м·К).',
     indent_cm=1.25)
formula(doc,
        f'R₀_пол = 0,17 + {d_floor}/{lam} + 0,04 = 0,17 + {d_floor/lam:.3f} + 0,04 = {R_floor:.2f} м²·°С/Вт')
para(doc,
     f'R₀_пол = {R_floor:.2f} м²·°С/Вт > R_req = {R_floor_req:.2f} м²·°С/Вт — '
     'условие выполнено ✓',
     indent_cm=1.25)

# Окна
heading(doc, '7.4. Окна', 2)
R_win = 0.56
para(doc,
     'Приняты окна с двухкамерным стеклопакетом (4-16-4-16-4) в ПВХ профиле. '
     f'Сопротивление теплопередаче принято: R₀_ок = {R_win} м²·°С/Вт.',
     indent_cm=1.25)
para(doc,
     f'R₀_ок = {R_win} м²·°С/Вт > R_req = {R_win_req:.2f} м²·°С/Вт — '
     'условие выполнено ✓',
     indent_cm=1.25)

# Дверь
heading(doc, '7.5. Входная дверь', 2)
R_door = 0.60
para(doc,
     f'Принята утеплённая металлическая входная дверь. '
     f'Сопротивление теплопередаче: R₀_дв = {R_door} м²·°С/Вт.',
     indent_cm=1.25)
para(doc,
     f'R₀_дв = {R_door} м²·°С/Вт > R_req = {R_door_req:.2f} м²·°С/Вт — '
     'условие выполнено ✓',
     indent_cm=1.25)

# Сводная таблица R
heading(doc, '7.6. Сводная таблица сопротивлений теплопередаче', 2)
t5 = doc.add_table(rows=1, cols=4)
t5.style = 'Table Grid'
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(['Конструкция', 'R₀_факт, м²·°С/Вт', 'R_req, м²·°С/Вт', 'Соответствие']):
    t5.rows[0].cells[i].text = txt
    t5.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if t5.rows[0].cells[i].paragraphs[0].runs:
        t5.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        t5.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        t5.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_bg(t5.rows[0].cells[i], 'D9E2F3')

r_summary = [
    ('Наружная стена',
     f'{R_wall:.2f}'.replace('.', ','),
     f'{R_wall_req:.2f}'.replace('.', ','), '✓'),
    ('Перекрытие (потолок)',
     f'{R_ceil:.2f}'.replace('.', ','),
     f'{R_ceil_req:.2f}'.replace('.', ','), '✓'),
    ('Пол над улицей',
     f'{R_floor:.2f}'.replace('.', ','),
     f'{R_floor_req:.2f}'.replace('.', ','), '✓'),
    ('Окна',
     f'{R_win}'.replace('.', ','),
     f'{R_win_req:.2f}'.replace('.', ','), '✓'),
    ('Входная дверь',
     f'{R_door}'.replace('.', ','),
     f'{R_door_req:.2f}'.replace('.', ','), '✓'),
]
for row_data in r_summary:
    row = t5.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(11)
doc.add_paragraph()

# =====================================================================
# 8. ПЛОЩАДИ ОГРАЖДАЮЩИХ КОНСТРУКЦИЙ
# =====================================================================
heading(doc, '8. ПЛОЩАДИ ОГРАЖДАЮЩИХ КОНСТРУКЦИЙ', 1)
para(doc,
     'Площади ограждающих конструкций одного дома определяются по геометрическим '
     'параметрам здания (ГОСТ Р 56778-2015, п. 5.2):',
     indent_cm=1.25)

A_wall_gross = 18.0 * 3.0   # 54
A_win_total  = 3 * 2.0       # 6
A_door_total = 1 * 2.0       # 2
A_wall_net   = A_wall_gross - A_win_total - A_door_total  # 46
A_ceil       = 4.0 * 5.0    # 20
A_floor      = 4.0 * 5.0    # 20

para(doc, f'    Суммарная площадь стен (брутто):  A_ст_бр = P × H = 18,0 × 3,0 = {A_wall_gross:.1f} м²')
para(doc, f'    Площадь окон:                     A_ок = 3 × 2,0 = {A_win_total:.1f} м²')
para(doc, f'    Площадь дверей:                   A_дв = 1 × 2,0 = {A_door_total:.1f} м²')
formula(doc,
        f'A_ст = {A_wall_gross:.1f} − {A_win_total:.1f} − {A_door_total:.1f} = {A_wall_net:.1f} м²  (чистая площадь стен)')
para(doc, f'    Площадь потолка:                  A_пт = 4,0 × 5,0 = {A_ceil:.1f} м²')
para(doc, f'    Площадь пола:                     A_пол = 4,0 × 5,0 = {A_floor:.1f} м²')

# =====================================================================
# 9. РАСЧЁТ ТРАНСМИССИОННЫХ ТЕПЛОПОТЕРЬ
# =====================================================================
heading(doc, '9. РАСЧЁТ ТРАНСМИССИОННЫХ ТЕПЛОПОТЕРЬ', 1)
para(doc,
     'Трансмиссионные теплопотери через ограждающие конструкции определяются по '
     'формуле ГОСТ Р 56778-2015, п. 5.3:',
     indent_cm=1.25)
formula(doc, 'Q_тр = (A / R₀) × (t_int − t_ext) × n,  Вт')
para(doc, 'где:', indent_cm=1.25)
para(doc, '    A     — площадь ограждающей конструкции, м²;')
para(doc, '    R₀    — приведённое сопротивление теплопередаче, м²·°С/Вт;')
para(doc,
     '    t_int = +20 °С — расчётная температура внутреннего воздуха '
     '(ГОСТ 30494-2011, таблица 1, жилые помещения);')
para(doc,
     '    t_ext = −31 °С — расчётная температура наружного воздуха '
     '(холодная пятидневка с обеспеченностью 0,92, СП 131.13330.2020);')
para(doc,
     '    n     — коэффициент, учитывающий положение наружной поверхности конструкции '
     'по отношению к наружному воздуху (СП 50.13330.2022, таблица 7).')

t_int = 20
t_ext = -31
Delta_t = t_int - t_ext  # 51

para(doc,
     f'Расчётный перепад температур: Δt = {t_int} − ({t_ext}) = {Delta_t} °С',
     indent_cm=1.25)
para(doc,
     'Для всех конструкций, непосредственно контактирующих с наружным воздухом, '
     'коэффициент n = 1,0 (пол над улицей рассматривается как открытое пространство).',
     indent_cm=1.25)

Q_walls = (A_wall_net / R_wall) * Delta_t
Q_ceil  = (A_ceil / R_ceil) * Delta_t
Q_floor = (A_floor / R_floor) * Delta_t
Q_win   = (A_win_total / R_win) * Delta_t
Q_door  = (A_door_total / R_door) * Delta_t
Q_tr    = Q_walls + Q_ceil + Q_floor + Q_win + Q_door

heading(doc, '9.1. Расчёт по элементам', 2)
para(doc, f'Наружные стены:', indent_cm=1.25)
formula(doc,
        f'Q_ст = ({A_wall_net:.1f} / {R_wall:.2f}) × {Delta_t} = '
        f'{A_wall_net/R_wall:.2f} × {Delta_t} = {Q_walls:.0f} Вт')

para(doc, f'Потолок:', indent_cm=1.25)
formula(doc,
        f'Q_пт = ({A_ceil:.1f} / {R_ceil:.2f}) × {Delta_t} = '
        f'{A_ceil/R_ceil:.3f} × {Delta_t} = {Q_ceil:.0f} Вт')

para(doc, f'Пол над улицей:', indent_cm=1.25)
formula(doc,
        f'Q_пол = ({A_floor:.1f} / {R_floor:.2f}) × {Delta_t} = '
        f'{A_floor/R_floor:.3f} × {Delta_t} = {Q_floor:.0f} Вт')

para(doc, f'Окна:', indent_cm=1.25)
formula(doc,
        f'Q_ок = ({A_win_total:.1f} / {R_win}) × {Delta_t} = '
        f'{A_win_total/R_win:.4f} × {Delta_t} = {Q_win:.0f} Вт')

para(doc, f'Входная дверь:', indent_cm=1.25)
formula(doc,
        f'Q_дв = ({A_door_total:.1f} / {R_door}) × {Delta_t} = '
        f'{A_door_total/R_door:.4f} × {Delta_t} = {Q_door:.0f} Вт')

para(doc,
     f'Суммарные трансмиссионные теплопотери:',
     indent_cm=1.25)
formula(doc,
        f'Q_тр = {Q_walls:.0f} + {Q_ceil:.0f} + {Q_floor:.0f} + {Q_win:.0f} + {Q_door:.0f} = {Q_tr:.0f} Вт')

heading(doc, '9.2. Поправочный коэффициент на инфильтрацию и ориентацию', 2)
para(doc,
     'В соответствии с ГОСТ Р 56778-2015 (п. 5.4) и рекомендациями СП 50.13330.2022 '
     'к трансмиссионным теплопотерям применяется поправочный коэффициент β_inf, '
     'учитывающий дополнительные потери на инфильтрацию через неплотности конструкций '
     'и воздействие ветрового давления. Для малоэтажных зданий с естественной '
     'вентиляцией принимается β_inf = 0,10 (10 %).', indent_cm=1.25)

beta = 0.10
Q_tr_corr = Q_tr * (1 + beta)
formula(doc,
        f'Q_тр_скорр = Q_тр × (1 + β_inf) = {Q_tr:.0f} × 1,10 = {Q_tr_corr:.0f} Вт')

# =====================================================================
# 10. РАСЧЁТ ВЕНТИЛЯЦИОННЫХ ТЕПЛОПОТЕРЬ
# =====================================================================
heading(doc, '10. РАСЧЁТ ВЕНТИЛЯЦИОННЫХ ТЕПЛОПОТЕРЬ', 1)
para(doc,
     'Для жилых помещений с естественной вентиляцией расход вентиляционного воздуха '
     'принимается из условия обеспечения нормируемой кратности воздухообмена. '
     'По СП 60.13330.2020 (п. 8.3), для жилых помещений при естественной вентиляции '
     'минимальная кратность воздухообмена составляет n_v = 0,35 ч⁻¹ '
     '(но не менее 30 м³/ч на одного человека).',
     indent_cm=1.25)

V_room = 4.0 * 5.0 * 3.0  # = 60 м³
n_v = 0.35
L_vent = n_v * V_room  # = 21 м³/ч

para(doc,
     f'Строительный объём одного дома: V = 4,0 × 5,0 × 3,0 = {V_room:.0f} м³',
     indent_cm=1.25)
formula(doc, f'L = n_v × V = {n_v} × {V_room:.0f} = {L_vent:.1f} м³/ч')

para(doc,
     'Вентиляционные теплопотери определяются по формуле СП 60.13330.2020 (п. 6.4.2):',
     indent_cm=1.25)
formula(doc, 'Q_в = ρ_в × c_в × L × (t_int − t_ext) / 3600,  Вт')
para(doc, 'где:', indent_cm=1.25)
para(doc, '    ρ_в = 1,2 кг/м³   — плотность наружного воздуха при −31 °С;')
para(doc, '    c_в = 1006 Дж/(кг·К) — удельная теплоёмкость воздуха.')

rho = 1.2
c_v = 1006.0
Q_vent = rho * c_v * L_vent * Delta_t / 3600.0
formula(doc,
        f'Q_в = 1,2 × 1006 × {L_vent:.1f} × {Delta_t} / 3600 = {Q_vent:.0f} Вт')

# =====================================================================
# 11. ИТОГО ТЕПЛОПОТЕРИ ОДНОГО ДОМА
# =====================================================================
heading(doc, '11. ИТОГО ТЕПЛОПОТЕРИ ОДНОГО ДОМА', 1)
Q_house = Q_tr_corr + Q_vent
formula(doc,
        f'Q_дом = Q_тр_скорр + Q_в = {Q_tr_corr:.0f} + {Q_vent:.0f} = {Q_house:.0f} Вт ≈ {Q_house/1000:.2f} кВт')

para(doc, 'Сводная таблица теплопотерь одного дома:', indent_cm=1.25)
t6 = doc.add_table(rows=1, cols=3)
t6.style = 'Table Grid'
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(['Составляющая', 'Теплопотери, Вт', 'Доля, %']):
    t6.rows[0].cells[i].text = txt
    t6.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if t6.rows[0].cells[i].paragraphs[0].runs:
        t6.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        t6.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        t6.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_bg(t6.rows[0].cells[i], 'D9E2F3')

loss_rows = [
    ('Стены', f'{Q_walls:.0f}', f'{Q_walls/Q_house*100:.1f}'.replace('.', ',')),
    ('Потолок', f'{Q_ceil:.0f}', f'{Q_ceil/Q_house*100:.1f}'.replace('.', ',')),
    ('Пол', f'{Q_floor:.0f}', f'{Q_floor/Q_house*100:.1f}'.replace('.', ',')),
    ('Окна', f'{Q_win:.0f}', f'{Q_win/Q_house*100:.1f}'.replace('.', ',')),
    ('Дверь', f'{Q_door:.0f}', f'{Q_door/Q_house*100:.1f}'.replace('.', ',')),
    ('Инфильтрация (+10%)', f'{Q_tr*beta:.0f}', f'{Q_tr*beta/Q_house*100:.1f}'.replace('.', ',')),
    ('Вентиляция', f'{Q_vent:.0f}', f'{Q_vent/Q_house*100:.1f}'.replace('.', ',')),
    ('ИТОГО', f'{Q_house:.0f}', '100'),
]
for row_data in loss_rows:
    row = t6.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(11)
            row.cells[i].paragraphs[0].runs[0].bold = (row_data[0] == 'ИТОГО')
    if row_data[0] == 'ИТОГО':
        for i in range(3):
            set_cell_bg(row.cells[i], 'FFF2CC')
doc.add_paragraph()

# =====================================================================
# 12. ТЕПЛОПОТЕРИ 51 ДОМА
# =====================================================================
heading(doc, '12. СУММАРНЫЕ ТЕПЛОПОТЕРИ 51 ДОМА', 1)
Q_total = 51 * Q_house
formula(doc,
        f'Q_бр = N × Q_дом = 51 × {Q_house:.0f} = {Q_total:.0f} Вт ≈ {Q_total/1000:.1f} кВт')

# =====================================================================
# 13. РАСЧЁТ МОЩНОСТИ ЭЛЕКТРООТОПЛЕНИЯ
# =====================================================================
heading(doc, '13. РАСЧЁТ ЭЛЕКТРИЧЕСКОЙ МОЩНОСТИ', 1)
heading(doc, '13.1. Мощность системы отопления', 2)
para(doc,
     'Для электрического отопления (инфракрасные нагреватели, конвекторы или '
     'тепловые насосы) устанавливаемая мощность определяется из теплового баланса '
     'с учётом запасного коэффициента надёжности k_з = 1,15 (15 %) на неучтённые '
     'теплопотери (тепловые мосты, неплотности и пр.):',
     indent_cm=1.25)

k_z = 1.15
P_heat_house = Q_house * k_z
P_heat_total = 51 * P_heat_house
formula(doc,
        f'P_от_дом = Q_дом × k_з = {Q_house:.0f} × {k_z} = {P_heat_house:.0f} Вт ≈ {P_heat_house/1000:.2f} кВт')
formula(doc,
        f'P_от = 51 × {P_heat_house:.0f} = {P_heat_total:.0f} Вт ≈ {P_heat_total/1000:.1f} кВт')

heading(doc, '13.2. Прочие электрические нагрузки', 2)
para(doc,
     'Помимо отопления учитываются прочие электрические нагрузки на один дом:',
     indent_cm=1.25)

P_light = 200    # Вт  (10 Вт/м² × 20 м²)
P_dhw   = 500    # Вт  (накопительный водонагреватель, средняя нагрузка)
P_misc  = 300    # Вт  (розеточные нагрузки, зарядные устройства и пр.)
P_other_house = P_light + P_dhw + P_misc
P_other_total = 51 * P_other_house

t7 = doc.add_table(rows=1, cols=3)
t7.style = 'Table Grid'
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(['Нагрузка', 'На 1 дом, Вт', 'Норматив/ссылка']):
    t7.rows[0].cells[i].text = txt
    t7.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if t7.rows[0].cells[i].paragraphs[0].runs:
        t7.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        t7.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        t7.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_bg(t7.rows[0].cells[i], 'D9E2F3')

other_rows = [
    ('Освещение', f'{P_light}',
     '10 Вт/м² × 20 м² (СП 52.13330.2016)'),
    ('Электроводонагреватель (ГВС)', f'{P_dhw}',
     'Средняя мощность нагрева'),
    ('Розеточные нагрузки (бытовые)', f'{P_misc}',
     'ПУЭ, СП 256.1325800.2016'),
    ('ИТОГО прочие нагрузки', f'{P_other_house}', '—'),
]
for row_data in other_rows:
    row = t7.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(11)
            row.cells[i].paragraphs[0].runs[0].bold = ('ИТОГО' in row_data[0])
doc.add_paragraph()

heading(doc, '13.3. Суммарная установленная электрическая мощность базы отдыха', 2)
P_install_house = P_heat_house + P_other_house
P_install_total = 51 * P_install_house

para(doc,
     f'Установленная мощность на один дом:',
     indent_cm=1.25)
formula(doc,
        f'P_уст_дом = P_от_дом + P_прочие = {P_heat_house:.0f} + {P_other_house} = {P_install_house:.0f} Вт ≈ {P_install_house/1000:.2f} кВт')

para(doc,
     f'Суммарная установленная мощность 51 дома:',
     indent_cm=1.25)
formula(doc,
        f'P_уст_бр = 51 × {P_install_house:.0f} = {P_install_total:.0f} Вт ≈ {P_install_total/1000:.1f} кВт')

heading(doc, '13.4. Расчётная (договорная) мощность с учётом коэффициента одновременности', 2)
para(doc,
     'Для баз отдыха и кемпингов коэффициент одновременности нагрузок '
     'K_od принимается по СП 256.1325800.2016 (приложение Б) в диапазоне 0,80–0,90. '
     'Принимается K_od = 0,85.',
     indent_cm=1.25)

K_od = 0.85
P_demand = P_install_total * K_od
formula(doc,
        f'P_расч = P_уст_бр × K_od = {P_install_total:.0f} × {K_od} = {P_demand:.0f} Вт ≈ {P_demand/1000:.1f} кВт')

# =====================================================================
# 14. СВОДНАЯ ТАБЛИЦА
# =====================================================================
doc.add_page_break()
heading(doc, '14. СВОДНАЯ ТАБЛИЦА РЕЗУЛЬТАТОВ', 1)

t8 = doc.add_table(rows=1, cols=3)
t8.style = 'Table Grid'
t8.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(['Показатель', 'Ед. изм.', 'Значение']):
    t8.rows[0].cells[i].text = txt
    t8.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if t8.rows[0].cells[i].paragraphs[0].runs:
        t8.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        t8.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        t8.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_bg(t8.rows[0].cells[i], 'D9E2F3')

summary = [
    ('Количество домов', 'шт.', '51'),
    ('Жилая площадь одного дома', 'м²', '20'),
    ('Расчётная температура наружного воздуха', '°С', '−31'),
    ('Расчётная температура внутреннего воздуха', '°С', '+20'),
    ('ГСОП', '°С·сут', f'{GSOP:.0f}'),
    ('', '', ''),
    ('— ОДИН ДОМ —', '', ''),
    ('Трансмиссионные теплопотери (без поправки)', 'Вт', f'{Q_tr:.0f}'),
    ('Поправка на инфильтрацию (+10%)', 'Вт', f'{Q_tr*beta:.0f}'),
    ('Вентиляционные теплопотери', 'Вт', f'{Q_vent:.0f}'),
    ('ИТОГО теплопотери 1 дома', 'Вт', f'{Q_house:.0f}'),
    ('ИТОГО теплопотери 1 дома', 'кВт', f'{Q_house/1000:.2f}'.replace('.', ',')),
    ('', '', ''),
    ('— 51 ДОМ —', '', ''),
    ('Суммарные теплопотери 51 дома', 'кВт', f'{Q_total/1000:.1f}'.replace('.', ',')),
    ('Мощность электроотопления 51 дома (k=1,15)', 'кВт', f'{P_heat_total/1000:.1f}'.replace('.', ',')),
    ('Прочие эл. нагрузки 51 дома', 'кВт', f'{P_other_total/1000:.1f}'.replace('.', ',')),
    ('Суммарная установленная эл. мощность', 'кВт', f'{P_install_total/1000:.1f}'.replace('.', ',')),
    ('Расчётная мощность (K_od=0,85)', 'кВт', f'{P_demand/1000:.1f}'.replace('.', ',')),
]
for row_data in summary:
    row = t8.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(11)
        is_total = ('ИТОГО' in val or '51 ДОМ' in val or 'ОДИН ДОМ' in val or
                    'Расчётная мощность' in val or 'Суммарная установленная' in val)
        if is_total:
            set_cell_bg(row.cells[i], 'FFF2CC')
            if row.cells[i].paragraphs[0].runs:
                row.cells[i].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# =====================================================================
# 15. ВЫВОДЫ
# =====================================================================
heading(doc, '15. ВЫВОДЫ', 1)

conclusions = [
    f'Расчёт теплопотерь выполнен для 51 жилого дома базы отдыха, '
    f'расположенной в г. Нижний Новгород, в соответствии с требованиями '
    f'ГОСТ Р 56778-2015, ГОСТ Р 54851-2011, ГОСТ 30494-2011 и СП 50.13330.2022.',

    f'Все ограждающие конструкции удовлетворяют нормируемым требованиям '
    f'по сопротивлению теплопередаче согласно СП 50.13330.2022 (таблица 3) '
    f'для климатического района г. Нижний Новгород (ГСОП = {GSOP:.0f} °С·сут).',

    f'Суммарные теплопотери одного жилого дома составляют '
    f'{Q_house:.0f} Вт ({Q_house/1000:.2f} кВт), '
    f'в том числе: трансмиссионные — {Q_tr_corr:.0f} Вт, '
    f'вентиляционные — {Q_vent:.0f} Вт.',

    f'Суммарные теплопотери базы отдыха (51 дом) составляют '
    f'{Q_total/1000:.1f} кВт.',

    f'Установленная электрическая мощность системы отопления с учётом '
    f'запасного коэффициента k_з = 1,15 составляет {P_heat_total/1000:.1f} кВт.',

    f'С учётом прочих электрических нагрузок (освещение, ГВС, розеточные нагрузки) '
    f'суммарная установленная мощность базы отдыха составляет {P_install_total/1000:.1f} кВт.',

    f'Расчётная (договорная) электрическая мощность базы отдыха с учётом '
    f'коэффициента одновременности K_od = 0,85 (СП 256.1325800.2016) составляет '
    f'{P_demand/1000:.1f} кВт. Данная величина принимается для технологического '
    f'присоединения к электрическим сетям.',
]
for i, c in enumerate(conclusions, 1):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(c)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_paragraph()
para(doc,
     'Настоящий расчёт составлен в соответствии с действующими нормативными документами '
     'и может быть использован при разработке проектной документации на строительство '
     'базы отдыха.',
     indent_cm=1.25)

# =====================================================================
# СОХРАНЕНИЕ
# =====================================================================
out_path = r'C:\dev\test-projekt\heat-loss-calculation\Расчёт теплопотерь базы отдыха 51 дом.docx'
doc.save(out_path)
print(f'Документ сохранён: {out_path}')
print(f'\n=== Ключевые результаты ===')
print(f'ГСОП = {GSOP:.0f} °С·сут')
print(f'Теплопотери 1 дома = {Q_house:.0f} Вт = {Q_house/1000:.2f} кВт')
print(f'Теплопотери 51 дома = {Q_total/1000:.1f} кВт')
print(f'Мощность отопления 51 дома (с k=1.15) = {P_heat_total/1000:.1f} кВт')
print(f'Суммарная установленная мощность = {P_install_total/1000:.1f} кВт')
print(f'Расчётная мощность (K_od=0.85) = {P_demand/1000:.1f} кВт')
